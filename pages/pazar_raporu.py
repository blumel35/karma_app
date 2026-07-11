"""
pages/pazar_raporu.py
─────────────────────
Karma App — Pazar Analiz Raporu
Seçilen kriterlere göre aktif + pasif ilanları çekip kapsamlı rapor üretir.
PDF + Word çıktısı.
"""

import streamlit as st
import sys, os, json, importlib, io, tempfile, pickle
from pathlib import Path
from datetime import datetime, date, timedelta
from dateutil.relativedelta import relativedelta
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "core"))

from core.auth import oturum_kontrol
from core.ui_helpers import render_navbar, render_page_header
from core.tuik_ilce_hacim import hacim_sinifi, hacim_rozet_rengi, ilce_satis_2025

if not oturum_kontrol():
    st.switch_page("pages/giris.py")

_k = st.session_state.get("kullanici", {})
render_navbar(
    user_role=_k.get("rol", "danisan"),
    user_name=_k.get("ad_soyad") or _k.get("ad", ""),
)

# ─────────────────────────────────────────
# ARŞİV (HAM VERİ) — her başarılı "Yeni Sorgu" ayrı bir kayıt olarak saklanır.
# NOT: Sayfadaki mevcut "Önceki Raporlarım" paneli sadece FİLTREYİ hatırlıyor
# (tekrar Revy'den çekmen gerekiyor). Bu yeni arşiv ise HAM VERİYİ de saklıyor,
# "Yükle" ile Revy'ye hiç gitmeden anında geri getirebiliyorsun.
# ─────────────────────────────────────────
RPR_GECMIS_KLASOR = ROOT / "revy_pazar_cikti" / "gecmis_rapor"
RPR_GECMIS_INDEX  = RPR_GECMIS_KLASOR / "index.json"

def rpr_gecmis_kaydet_yerel(df_aktif_ham, df_pasif_ham, meta: dict):
    try:
        RPR_GECMIS_KLASOR.mkdir(parents=True, exist_ok=True)
        zaman = datetime.now()
        dosya_adi = f"rpr_{zaman.strftime('%Y%m%d_%H%M%S')}.pkl"
        with open(RPR_GECMIS_KLASOR / dosya_adi, "wb") as f:
            pickle.dump({"aktif": df_aktif_ham, "pasif": df_pasif_ham, "meta": meta}, f)
        index = []
        if RPR_GECMIS_INDEX.exists():
            try:
                index = json.loads(RPR_GECMIS_INDEX.read_text(encoding="utf-8"))
            except Exception:
                index = []
        index.insert(0, {
            "dosya": dosya_adi,
            "zaman": zaman.strftime("%d.%m.%Y %H:%M"),
            "etiket": meta.get("baslik", ""),
            "aktif_n": len(df_aktif_ham),
            "pasif_n": len(df_pasif_ham),
        })
        RPR_GECMIS_INDEX.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass

def rpr_gecmis_listele_yerel():
    if not RPR_GECMIS_INDEX.exists():
        return []
    try:
        return json.loads(RPR_GECMIS_INDEX.read_text(encoding="utf-8"))
    except Exception:
        return []

def rpr_gecmis_yukle_yerel(dosya_adi: str):
    try:
        with open(RPR_GECMIS_KLASOR / dosya_adi, "rb") as f:
            return pickle.load(f)
    except Exception:
        return None

def rpr_gecmis_sil_yerel(dosya_adi: str):
    try:
        (RPR_GECMIS_KLASOR / dosya_adi).unlink(missing_ok=True)
        index = [k for k in rpr_gecmis_listele_yerel() if k["dosya"] != dosya_adi]
        RPR_GECMIS_INDEX.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass

baslik_col, sorgu_col = st.columns([4, 1.4])
with baslik_col:
    render_page_header(
        "Pazar Analiz Raporu",
        "Seçilen bölge ve kriterlere göre kapsamlı pazar analizi ve PDF/Word raporu"
    )
with sorgu_col:
    st.write("")
    analiz_btn = st.button("🆕 Yeni Sorgu\n(Revy'den Çek)", type="primary", use_container_width=True, key="rpr_analiz",
        help="Revy'den TAZE veri çeker. Mahalle ve Mülk Türü filtreleri zaten çekilmiş veri üzerinde anında uygulanır.")

with st.expander("📚 Geçmiş Taramalar (Ham Veri Arşivi)", expanded=False):
    st.caption(
        "Her başarılı 'Yeni Sorgu' taraması burada HAM VERİSİYLE saklanır. "
        "'Yükle' ile Revy'ye hiç gitmeden anında geri getirebilirsin."
    )
    _rpr_gecmis_yerel = rpr_gecmis_listele_yerel()
    if not _rpr_gecmis_yerel:
        st.info("Henüz arşivlenmiş bir tarama yok.")
    else:
        st.caption(f"{len(_rpr_gecmis_yerel)} kayıtlı tarama")
        for _rg in _rpr_gecmis_yerel:
            _rc1, _rc2, _rc3, _rc4, _rc5 = st.columns([2.5, 1.3, 1.2, 1, 1])
            with _rc1:
                st.markdown(f'<div style="font-size:13px;font-weight:700;color:#0f172a;padding:6px 0;">{_rg.get("etiket","")}</div>', unsafe_allow_html=True)
            with _rc2:
                st.markdown(f'<div style="font-size:11px;color:#64748b;padding:6px 0;">🕐 {_rg.get("zaman","-")}</div>', unsafe_allow_html=True)
            with _rc3:
                st.markdown(f'<div style="font-size:12px;color:#355C7D;font-weight:700;padding:6px 0;">Aktif:{_rg.get("aktif_n",0)} Pasif:{_rg.get("pasif_n",0)}</div>', unsafe_allow_html=True)
            with _rc4:
                if st.button("📂 Yükle", key=f"rpr_gecmis_yukle_yerel_{_rg['dosya']}", use_container_width=True):
                    _paket = rpr_gecmis_yukle_yerel(_rg["dosya"])
                    if _paket is not None:
                        st.session_state["rpr_aktif_ham"] = _paket["aktif"]
                        st.session_state["rpr_pasif_ham"] = _paket["pasif"]
                        st.session_state["rpr_meta"] = _paket["meta"]
                        st.toast(f"'{_rg.get('etiket','')}' yüklendi ✓")
                        st.rerun()
                    else:
                        st.error("Bu kayıt yüklenemedi.")
            with _rc5:
                if st.button("🗑", key=f"rpr_gecmis_sil_yerel_{_rg['dosya']}", use_container_width=True, help="Bu arşiv kaydını sil"):
                    rpr_gecmis_sil_yerel(_rg["dosya"])
                    st.toast("Kayıt silindi.")
                    st.rerun()
            st.divider()

# ─── YARDIMCI FONKSİYONLAR ────────────────────────────────────────────────────
def _fmt(v, suffix="₺", decimals=0):
    if v is None or (isinstance(v, float) and (np.isnan(v) or np.isinf(v))): return "-"
    fmt = f"{{:,.{decimals}f}} {suffix}".format(v).replace(",",".")
    return fmt

def parse_num(v):
    try: return float(str(v).replace(".","").replace(",",".").replace("₺","").replace("TL","").strip())
    except: return None

IZMIR_ILCELER = {
    "Aliağa":487,"Balçova":488,"Bayındır":489,"Bayraklı":490,
    "Bergama":491,"Beydağ":492,"Bornova":493,"Buca":494,
    "Çeşme":495,"Çiğli":496,"Dikili":497,"Foça":498,
    "Gaziemir":499,"Güzelbahçe":500,"Karabağlar":501,"Karaburun":502,
    "Karşıyaka":503,"Kemalpaşa":504,"Kınık":505,"Kiraz":506,
    "Konak":507,"Menderes":508,"Menemen":509,"Narlıdere":510,
    "Ödemiş":511,"Seferihisar":512,"Selçuk":513,"Tire":514,
    "Torbalı":515,"Urla":516,
}
MULK_TIPLERI  = {"Konut":"1","Ticari":"2","Arsa":"3"}
ISLEM_TIPLERI = {"Satılık":"sale","Kiralık":"rent"}
MULK_TURU_OPTS = [
    "Daire","Apartman Dairesi","Konut","Toplu Konut","Ticari Konut",
    "Villa","Müstakil Ev","Bina","Komple Bina",
    "Dükkan & Mağaza","Büro & Ofis","Depo & Antrepo",
    "Fabrika","Büfe","Restoran & Lokanta","Pazar Yeri",
    "Tarla","Arazi","Bağ","Zeytinlik","Diğer",
]

@st.cache_data
def mahalle_yukle():
    yol = ROOT / "core" / "izmir_mahalleler.json"
    if yol.exists():
        return json.load(open(yol, encoding="utf-8"))
    return {}

@st.cache_data
def ofis_duzeltme_yukle():
    yol = ROOT / "core" / "ofis_duzeltme.json"
    if yol.exists():
        return json.load(open(yol, encoding="utf-8"))
    return {}

MAHALLELER = mahalle_yukle()
OFIS_DUZELTME = ofis_duzeltme_yukle()

import re as _re
def ad_normalize(v):
    if pd.isna(v): return None
    s = str(v).strip()
    if not s: return None
    s = s.replace("İ","I").replace("ı","i")
    s = _re.sub(r"\s+", " ", s)
    return s.upper()

def ofis_duzelt(v):
    n = ad_normalize(v)
    if n is None: return None
    return OFIS_DUZELTME.get(n, n)

def revy_kimlik_al():
    try:
        rv = st.secrets.get("revy", {})
        if rv.get("kullanici") and rv.get("sifre"):
            return {"revy1_kullanici": rv["kullanici"], "revy1_sifre": rv["sifre"],
                    "revy_giris_url": rv.get("giris_url","https://revy.com.tr")}
    except Exception: pass
    try:
        import core.revy_pazar_cek as rpc
        return rpc.ayarlari_oku(ROOT / "ayarlar.txt")
    except Exception: pass
    return None

# ─── SUPABASE YARDIMCILARI ─────────────────────────────────────────────────────
def _supa():
    try:
        from supabase import create_client
        url = st.secrets.get("SUPABASE_URL","") or st.secrets.get("supabase",{}).get("url","")
        key = st.secrets.get("SUPABASE_KEY","") or st.secrets.get("supabase",{}).get("secret_key","") or st.secrets.get("supabase",{}).get("publishable_key","")
        if url and key:
            return create_client(url, key)
    except Exception: pass
    return None

def rapor_gecmis_kaydet(meta: dict, aktif_n: int, pasif_n: int):
    supa = _supa()
    if not supa: return
    uid = _k.get("id","") or _k.get("user_key","")
    try:
        supa.table("pazar_rapor_gecmis").upsert({
            "kullanici_id": uid,
            "baslik": meta.get("baslik",""),
            "filtre_json": json.dumps(meta, ensure_ascii=False, default=str),
            "aktif_n": aktif_n,
            "pasif_n": pasif_n,
            "tarih": datetime.now().isoformat(),
        }).execute()
    except Exception: pass

def rapor_gecmis_yukle() -> list:
    supa = _supa()
    if not supa: return []
    uid = _k.get("id","") or _k.get("user_key","")
    try:
        r = supa.table("pazar_rapor_gecmis").select("*").eq("kullanici_id", uid).order("tarih", desc=True).limit(10).execute()
        return r.data or []
    except Exception: return []

def rapor_gecmis_yukle_filtre(kayit: dict):
    try:
        f = json.loads(kayit["filtre_json"])
        st.session_state["rpr_ilce"]  = f.get("ilce", [])
        st.session_state["rpr_mah"]   = f.get("mah", [])
        st.session_state["rpr_islem"] = f.get("islem", ["Satılık"])
        st.session_state["rpr_mulk"]  = f.get("mulk", ["Konut"])
        st.session_state["rpr_tur"]   = f.get("tur", [])
        donem_ters = {30:"Son 1 ay",90:"Son 3 ay",180:"Son 6 ay",365:"Son 12 ay",548:"Son 18 ay",730:"Son 24 ay"}
        st.session_state["rpr_donem"] = donem_ters.get(f.get("gun",365), "Son 12 ay")
    except Exception: pass

# ─── GEÇMİŞ ARAMALAR PANELİ ──────────────────────────────────────────────────
_gecmis = rapor_gecmis_yukle()
if _gecmis:
    with st.expander(f"📋 Önceki Raporlarım ({len(_gecmis)})", expanded=False):
        for _g in _gecmis:
            gc1, gc2, gc3 = st.columns([4, 1.5, 1])
            with gc1:
                st.markdown(f'<div style="font-size:13px;font-weight:600;padding:4px 0;">{_g["baslik"]}</div>', unsafe_allow_html=True)
                try:
                    _t = datetime.fromisoformat(_g["tarih"]).strftime("%d.%m.%Y %H:%M")
                except: _t = ""
                st.caption(f"🕐 {_t} · Aktif: {_g.get('aktif_n',0)} · Pasif: {_g.get('pasif_n',0)}")
            with gc2:
                if st.button("📂 Yükle", key=f"rpr_yukle_{_g['id']}", use_container_width=True):
                    rapor_gecmis_yukle_filtre(_g)
                    st.toast("Filtre yüklendi — Yeni Sorgu'ya basın ✓")
                    st.rerun()
            with gc3:
                if st.button("🗑", key=f"rpr_sil_{_g['id']}", use_container_width=True):
                    try:
                        _supa().table("pazar_rapor_gecmis").delete().eq("id", _g["id"]).execute()
                        st.rerun()
                    except Exception: pass
            st.divider()

# ─── FİLTRE PANELİ ────────────────────────────────────────────────────────────
with st.expander("🔍 Rapor Filtresi", expanded=True):
    st.markdown('<div style="font-size:10px;font-weight:800;color:#1e2d3d;text-transform:uppercase;letter-spacing:.08em;margin-bottom:8px;">📍 Konum</div>', unsafe_allow_html=True)
    rc1, rc2 = st.columns([2, 3])
    with rc1:
        sec_ilce = st.multiselect("İlçe", list(IZMIR_ILCELER.keys()), key="rpr_ilce")
    with rc2:
        mah_opts = sorted(set(m for i in sec_ilce for m in MAHALLELER.get(i, [])))
        sec_mah = st.multiselect("Mahalle", mah_opts, placeholder="Tümü (opsiyonel)", key="rpr_mah")

    st.markdown('<hr style="border:none;border-top:1px solid #e2e8f0;margin:12px 0;">', unsafe_allow_html=True)
    st.markdown('<div style="font-size:10px;font-weight:800;color:#1e2d3d;text-transform:uppercase;letter-spacing:.08em;margin-bottom:8px;">📋 İlan Özellikleri</div>', unsafe_allow_html=True)
    rc3, rc4, rc5 = st.columns([1.5, 1.5, 2])
    with rc3:
        sec_islem = st.multiselect("İşlem Tipi", list(ISLEM_TIPLERI.keys()),
            default=["Satılık"], key="rpr_islem")
    with rc4:
        sec_mulk = st.multiselect("Mülk Tipi", list(MULK_TIPLERI.keys()),
            default=["Konut"], key="rpr_mulk")
    with rc5:
        sec_tur = st.multiselect("Mülk Türü", MULK_TURU_OPTS,
            placeholder="Tümü", key="rpr_tur")

    st.markdown('<hr style="border:none;border-top:1px solid #e2e8f0;margin:12px 0;">', unsafe_allow_html=True)
    st.markdown('<div style="font-size:10px;font-weight:800;color:#1e2d3d;text-transform:uppercase;letter-spacing:.08em;margin-bottom:8px;">📅 Analiz Dönemi (Pasif İlanlar)</div>', unsafe_allow_html=True)
    pasif_gun = st.selectbox("İlandan Kalkma Dönemi",
        ["Son 1 ay","Son 3 ay","Son 6 ay","Son 12 ay","Son 18 ay","Son 24 ay"],
        index=3, key="rpr_donem")

btn_col1, btn_col2 = st.columns([3, 1])
with btn_col2:
    filtrele_btn = st.button("🔍 Filtrele", type="primary",
        use_container_width=True, key="rpr_filtrele_btn",
        help="Revy'ye gitmez — zaten çekilmiş ham veri üzerinde Mahalle/Mülk Türü filtresini anında uygular.")

# ─── VERİ ÇEKME ────────────────────────────────────────────────────────────────
DONEM_GUN = {
    "Son 1 ay": 30, "Son 3 ay": 90, "Son 6 ay": 180,
    "Son 12 ay": 365, "Son 18 ay": 548, "Son 24 ay": 730
}

if analiz_btn:
    if not sec_ilce:
        st.error("En az bir ilçe seçin.")
        st.stop()
    if not sec_islem or not sec_mulk:
        st.error("İşlem tipi ve mülk tipi zorunlu.")
        st.stop()

    ayarlar = revy_kimlik_al()
    if not ayarlar:
        st.error("❌ Revy hesabı bulunamadı. ayarlar.txt veya secrets.toml kontrol edin.")
        st.stop()

    durum_ph = st.empty()
    pb = st.progress(0, text="Hazırlanıyor...")

    def pcb(msg): durum_ph.info("🔄 Veri çekiliyor, lütfen bekleyiniz...")

    try:
        import core.revy_pazar_cek as rpc
        importlib.reload(rpc)

        pcb("Revy'ye bağlanılıyor...")
        cookies = rpc.selenium_cookie_al(
            kullanici=ayarlar["revy1_kullanici"],
            sifre=ayarlar["revy1_sifre"],
            giris_url=ayarlar.get("revy_giris_url","https://revy.com.tr"),
            headless=True, progress_cb=pcb)

        if not cookies:
            durum_ph.error("❌ Revy cookie alınamadı — kullanıcı adı/şifre kontrol edin.")
            st.stop()

        pb.progress(25, text="Aktif ilanlar çekiliyor...")

        _islem_map = {"Satılık": "satilik", "Kiralık": "kiralik"}
        _mulk_lower = [m.lower() for m in sec_mulk]
        _islem_lower = [_islem_map.get(i, i.lower()) for i in sec_islem]

        filtre_aktif = {
            "ilce": sec_ilce, "mulk": _mulk_lower,
            "islem": _islem_lower, "durum": ["aktif"],
        }
        cikti_aktif = rpc.pazar_cek(cookies, filtre_aktif,
            cikti_klasor=ROOT/"revy_pazar_cikti", progress_cb=pcb)

        pb.progress(60, text="Pasif ilanlar çekiliyor...")

        gun = DONEM_GUN.get(pasif_gun, 365)
        tampon_gun = gun + 365
        bas = (datetime.now() - timedelta(days=tampon_gun)).strftime("%Y-%m-%d")
        bit = datetime.now().strftime("%Y-%m-%d")
        filtre_pasif = {
            "ilce": sec_ilce, "mulk": _mulk_lower, "islem": _islem_lower,
            "durum": ["yayindan_kalkan"], "baslangic": bas, "bitis": bit,
        }
        cikti_pasif = rpc.pazar_cek(cookies, filtre_pasif,
            cikti_klasor=ROOT/"revy_pazar_cikti", progress_cb=pcb)

        pb.progress(90, text="Analiz hazırlanıyor...")

        df_aktif = cikti_aktif.get("aktif", pd.DataFrame())
        df_pasif = cikti_pasif.get("yayindan_kalkan", pd.DataFrame())

        if df_aktif.empty and df_pasif.empty:
            st.warning(f"⚠️ Debug: cikti_aktif keys={list(cikti_aktif.keys())}, cikti_pasif keys={list(cikti_pasif.keys())}")
            st.warning(f"Filtre aktif: {filtre_aktif}")
            st.warning(f"Filtre pasif: {filtre_pasif}")

        # NOT: Mahalle/Mülk Türü filtresi artık BURADA uygulanmıyor —
        # ham veri saklanıyor, filtre her rerun'da ayrıca uygulanıyor
        # (bkz. "RAPOR GÖRÜNTÜLEME" öncesi blok). Böylece "Filtrele"
        # butonu Revy'ye gitmeden anında çalışabiliyor.

        _baslik_parcalar = []
        if sec_ilce: _baslik_parcalar.append(" / ".join(sec_ilce))
        if sec_islem: _baslik_parcalar.append(" + ".join(sec_islem))
        if sec_mulk:  _baslik_parcalar.append(" + ".join(sec_mulk))
        _baslik_parcalar.append(pasif_gun)
        auto_baslik = " · ".join(_baslik_parcalar)

        st.session_state["rpr_aktif_ham"] = df_aktif
        st.session_state["rpr_pasif_ham"] = df_pasif
        st.session_state["rpr_meta"] = {
            "ilce": sec_ilce, "mah": sec_mah, "islem": sec_islem,
            "mulk": sec_mulk, "tur": sec_tur, "donem": pasif_gun,
            "baslik": auto_baslik,
            "tarih": datetime.now().strftime("%d.%m.%Y %H:%M"),
            "gun": gun,
        }

        rapor_gecmis_kaydet(st.session_state["rpr_meta"], len(df_aktif), len(df_pasif))
        rpr_gecmis_kaydet_yerel(df_aktif, df_pasif, st.session_state["rpr_meta"])
        pb.progress(100, text="Tamamlandı!")
        durum_ph.success(f"✅ Aktif: {len(df_aktif)} ilan | Pasif ({pasif_gun}): {len(df_pasif)} ilan")

    except Exception as e:
        durum_ph.error(f"Hata: {e}")
        pb.empty()

# ─── RAPOR GÖRÜNTÜLEME ────────────────────────────────────────────────────────
if "rpr_aktif_ham" not in st.session_state:
    st.info("👆 Filtreleri seçip **Yeni Sorgu** butonuna basın.")
    st.stop()

# Mahalle / Mülk Türü — HAM veri üzerinde HER ÇALIŞMADA uygulanan yerel filtre
# (Revy'ye gitmeden, "Filtrele" butonuna da ihtiyaç duymadan zaten anında
# uygulanır — buton sadece kullanıcıya görsel bir "uygulandı" hissi verir).
df_a = st.session_state["rpr_aktif_ham"].copy()
df_p = st.session_state["rpr_pasif_ham"].copy()

if sec_mah:
    if not df_a.empty and "Mahalle" in df_a.columns:
        df_a = df_a[df_a["Mahalle"].isin(sec_mah)]
    if not df_p.empty and "Mahalle" in df_p.columns:
        df_p = df_p[df_p["Mahalle"].isin(sec_mah)]

if sec_tur:
    if not df_a.empty and "Mülk türü" in df_a.columns:
        df_a = df_a[df_a["Mülk türü"].isin(sec_tur)]
    if not df_p.empty and "Mülk türü" in df_p.columns:
        df_p = df_p[df_p["Mülk türü"].isin(sec_tur)]

st.session_state["rpr_aktif"] = df_a
st.session_state["rpr_pasif"] = df_p
meta = st.session_state["rpr_meta"]

if df_a.empty and df_p.empty:
    st.warning("Seçilen kriterlere uygun ilan bulunamadı.")
    st.stop()

def hazirla(df, pasif=False):
    if df.empty:
        # Boş sonuç da geçerli bir durum (örn. aktif ilan yok) —
        # aşağıdaki rapor kodu __birim/__fiyat/__m2/__sure/__ofis
        # sütunlarını bulamazsa KeyError verir. 0 satırlı ama
        # doğru sütunlu bir df döndürerek bunu önlüyoruz.
        df = df.copy()
        df["__fiyat"] = pd.Series(dtype=float)
        df["__m2"]    = pd.Series(dtype=float)
        df["__birim"] = pd.Series(dtype=float)
        df["__sure"]  = pd.Series(dtype=float)
        df["__ofis"]  = pd.Series(dtype=str)
        df["__tarih"] = pd.Series(dtype="datetime64[ns]")
        df["__ay"]    = pd.Series(dtype=object)
        if pasif:
            df["__kalktarih"] = pd.Series(dtype="datetime64[ns]")
            df["__kalk_ay"]   = pd.Series(dtype=object)
        return df
    df = df.copy()
    df["__fiyat"] = df.get("Fiyat", pd.Series(dtype=float)).apply(parse_num)
    df["__m2"]    = pd.to_numeric(df.get("M2", pd.Series(dtype=float)), errors="coerce")
    df["__birim"] = df["__fiyat"] / df["__m2"].replace(0, pd.NA)
    df["__sure"]  = pd.to_numeric(df.get("İlan Yayın Süresi", pd.Series(dtype=float)), errors="coerce")
    df["__ofis"]  = df.get("Ofis", pd.Series(dtype=str)).apply(ofis_duzelt)
    if "İlan tarihi" in df.columns:
        df["__tarih"] = pd.to_datetime(df["İlan tarihi"], format="%d.%m.%Y", errors="coerce")
        df["__ay"] = df["__tarih"].dt.to_period("M")
    if pasif and "Yayından Kalkış Tarihi" in df.columns:
        df["__kalktarih"] = pd.to_datetime(df["Yayından Kalkış Tarihi"], errors="coerce")
        df["__kalk_ay"]   = df["__kalktarih"].dt.to_period("M")
    return df

def donem_basi_stok(aktif_df, pasif_df, tarih):
    n = 0
    if "__tarih" in aktif_df.columns:
        n += aktif_df[aktif_df["__tarih"].notna() & (aktif_df["__tarih"] < tarih)].shape[0]
    if "__tarih" in pasif_df.columns and "__kalktarih" in pasif_df.columns:
        mask = (
            pasif_df["__tarih"].notna() & pasif_df["__kalktarih"].notna() &
            (pasif_df["__tarih"] < tarih) & (pasif_df["__kalktarih"] >= tarih)
        )
        n += pasif_df[mask].shape[0]
    return n

df_a = hazirla(df_a, pasif=False)
df_p = hazirla(df_p, pasif=True)

baslik = meta["baslik"] or (
    " / ".join(meta["ilce"]) +
    (f" - {', '.join(meta['mah'])}" if meta["mah"] else "") +
    f" | {', '.join(meta['islem'])} | {', '.join(meta['mulk'])}"
)

st.markdown(f"""
<div style="background:linear-gradient(135deg,#1e2d3d 0%,#2a3f56 100%);
border-radius:12px;padding:20px 28px;margin-bottom:20px;color:white;">
  <div style="font-size:11px;font-weight:700;letter-spacing:.1em;color:#94a3b8;text-transform:uppercase;margin-bottom:4px;">
    PAZAR ANALİZ RAPORU</div>
  <div style="font-size:1.4rem;font-weight:800;color:white;">{baslik}</div>
  <div style="font-size:12px;color:#94a3b8;margin-top:6px;">
    📅 Rapor tarihi: {meta['tarih']} &nbsp;·&nbsp;
    📊 Aktif: {len(df_a)} ilan &nbsp;·&nbsp;
    📉 Pasif ({meta['donem']}): {len(df_p)} ilan
  </div>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# BÖLÜM 1: PAZAR ÖZETİ
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("### 1️⃣ Pazar Özeti")

gun = meta.get("gun", 365)
_sinir = pd.Timestamp.now() - pd.Timedelta(days=gun)
donem_bas_stok_n = donem_basi_stok(df_a, df_p, _sinir)

kapanan = 0
if "__kalktarih" in df_p.columns:
    kapanan = df_p[df_p["__kalktarih"].notna() & (df_p["__kalktarih"] >= _sinir)].shape[0]

absorp = (kapanan / donem_bas_stok_n * 100) if donem_bas_stok_n > 0 else 0

med_sure_p = df_p["__sure"].median() if not df_p.empty else None
med_sure_a = df_a["__sure"].median() if not df_a.empty else None

# ── TÜİK İLÇE HACİM SINIFI ──────────────────────────────────────────────────
# Bu, DUSUK_HACIM_ESIGI kontrolünden FARKLI bir şey: o, BU RAPORUN kendi
# örneklem büyüklüğüne (stok/kapanan sayısı) bakıyordu — mahalle bazlı
# filtrelemede bile küçük çıkabilir. Bu ise TÜİK'in resmi 2025 verisine göre
# İLÇENİN GENELİNDE pazar ne kadar büyük, onu gösteriyor — ikisi birlikte
# daha eksiksiz bir bağlam veriyor.
_ilce_hacim_sinifi_seti = set(hacim_sinifi(i) for i in meta["ilce"]) if meta.get("ilce") else set()
if len(_ilce_hacim_sinifi_seti) == 1:
    tuik_hacim = next(iter(_ilce_hacim_sinifi_seti))
elif len(_ilce_hacim_sinifi_seti) > 1:
    tuik_hacim = f"Karma ({len(meta['ilce'])} ilçe)"
else:
    tuik_hacim = "Bilinmiyor"

# ── DÜŞÜK HACİM KONTROLÜ ─────────────────────────────────────────────────────
# NOT: Dönem başı stok çok küçükse (ör. 5), yüzde (absorpsiyon) rakamı
# istatistiksel olarak anlamsız/gürültülü oluyor — 5 stoktan 13 kapanış
# %260 gibi aşırı bir sayı üretip yanlışlıkla "Hızlı Piyasa" etiketine yol
# açabiliyor, hele "şu an 0 aktif ilan" varken bu hiç doğru bir sinyal değil.
# Bu yüzden örneklem çok küçükse (stok VEYA kapanan sayısı eşik altındaysa)
# hız sınıflandırması yerine "Yetersiz Veri / Düşük Hacim" etiketi veriyoruz.
DUSUK_HACIM_ESIGI = 10
yetersiz_veri = (donem_bas_stok_n < DUSUK_HACIM_ESIGI) or (kapanan < DUSUK_HACIM_ESIGI)

if yetersiz_veri:
    pazar_hiz = "⚠️ Yetersiz Veri / Düşük Hacim"
    pazar_tip = (
        f"Bu bölgede ilan sirkülasyonu çok düşük (dönem başı stok: {donem_bas_stok_n}, "
        f"kapanan: {kapanan}) — yüzdesel pazar hızı bu örneklemde güvenilir değil, yorumlanmamalı"
    )
elif absorp > 70:
    pazar_hiz = "🔥 Hızlı Piyasa"
    pazar_tip = "Satıcı piyasası — ilanlar hızlı kapanıyor, fiyatlar yükseliş baskısında"
elif absorp > 45:
    pazar_hiz = "📈 Aktif Piyasa"
    pazar_tip = "Dengeli ama satıcı lehine — ilanlar makul sürede kapanıyor"
elif absorp > 25:
    pazar_hiz = "😐 Dengeli Piyasa"
    pazar_tip = "Arz ve talep dengeli — alıcılar seçenek buluyor, satıcılar bekliyor"
else:
    pazar_hiz = "❄️ Yavaş Piyasa"
    pazar_tip = "Alıcı piyasası — çok ilan var, az alıcı. Satıcılar fiyat indirimine açık"

# Hız + TÜİK hacim birleşik yorum (rapor metnine ek bağlam olarak eklenir)
if not yetersiz_veri and tuik_hacim not in ("Bilinmiyor", ""):
    if tuik_hacim in ("Yüksek Hacim", "Orta-Üst Hacim"):
        hacim_notu = f"Bu ilçe TÜİK 2025 verisine göre **{tuik_hacim.lower()}** bir pazar — bu oran güvenilir bir sinyal."
    else:
        hacim_notu = (
            f"Ancak bu ilçe TÜİK 2025 verisine göre **{tuik_hacim.lower()}** bir pazar — "
            f"yüzdesel oran doğru olsa da, ilçe genelinde işlem hacmi zaten düşük, temkinli yorumla."
        )
else:
    hacim_notu = ""

# ── PATCH 2: FSBO — kesin eşleşme ("Mülk Sahibi") ────────────────────────────
fsbo_a = (df_a["İlan sahibi türü"] == "Mülk Sahibi").sum() if "İlan sahibi türü" in df_a.columns else 0
fsbo_p = (df_p["İlan sahibi türü"] == "Mülk Sahibi").sum() if "İlan sahibi türü" in df_p.columns else 0
fsbo_oran = fsbo_a / len(df_a) * 100 if len(df_a) > 0 else 0

k = st.columns(6)
kart_data = [
    ("Şu An Piyasadaki İlan",   str(len(df_a)),                "aktif ilanlar",           "#1e2d3d"),
    (f"Kapanan İlan ({meta['donem']})", str(len(df_p)),        "satılan/yayından kalkan",  "#1e2d3d"),
    ("Pazar Hızı",              f"%{absorp:.0f}",               pazar_hiz,                 ("#64748b" if yetersiz_veri else ("#059669" if absorp > 45 else "#d97706"))),
    ("Medyan Satış Fiyatı/m²",  _fmt(df_a["__birim"].median()), "aktif ilanlar",           "#1e40af"),
    ("Acentesiz İlan (FSBO)",   f"{fsbo_a} aktif",             f"geçmişte: {fsbo_p} ilan · %{fsbo_oran:.0f} pazar payı", "#7c3aed"),
    ("İlçe Hacim Sınıfı (TÜİK)", tuik_hacim,                    "2025 resmi konut satış verisine göre", hacim_rozet_rengi(tuik_hacim)),
]
for col, (lbl, val, sub, color) in zip(k, kart_data):
    col.markdown(f"""<div style="background:white;border:1px solid #dce4ee;border-radius:10px;
        padding:14px 16px;box-shadow:0 2px 6px rgba(15,23,42,.04);height:100%;">
        <div style="font-size:9.5px;font-weight:700;color:#64748b;text-transform:uppercase;letter-spacing:.06em;margin-bottom:6px;">{lbl}</div>
        <div style="font-size:1.5rem;font-weight:800;color:{color};margin-bottom:4px;">{val}</div>
        <div style="font-size:10px;color:#94a3b8;line-height:1.4;">{sub}</div>
    </div>""", unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

k2 = st.columns(4)
kart_data2 = [
    ("Medyan Satış Fiyatı",         _fmt(df_a["__fiyat"].median()),  "şu an piyasadaki ilanların ortanca fiyatı"),
    ("Geçmiş Dönem Kapanış Fiyatı/m²", _fmt(df_p["__birim"].median()), f"son {meta['donem']} içinde kapanan ilanların ortanca m² fiyatı"),
    ("Ortalama Satış Süresi",        _fmt(med_sure_p, "gün"),         f"ilan açılıp kapanana kadar geçen süre (son {meta['donem']})"),
    ("Şu An Bekleyen İlan Süresi",   _fmt(med_sure_a, "gün"),         "aktif ilanların ortalama ne kadar süredir piyasada olduğu"),
]
for col, (lbl, val, sub) in zip(k2, kart_data2):
    col.markdown(f"""<div style="background:white;border:1px solid #dce4ee;border-radius:10px;
        padding:14px 16px;box-shadow:0 2px 6px rgba(15,23,42,.04);height:100%;">
        <div style="font-size:9.5px;font-weight:700;color:#64748b;text-transform:uppercase;letter-spacing:.06em;margin-bottom:6px;">{lbl}</div>
        <div style="font-size:1.3rem;font-weight:800;color:#1e2d3d;margin-bottom:4px;">{val}</div>
        <div style="font-size:10px;color:#94a3b8;line-height:1.4;">{sub}</div>
    </div>""", unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

med_fiyat_aktif = df_a["__birim"].median()
med_fiyat_pasif = df_p["__birim"].median()
fiyat_trend = ""
if pd.notna(med_fiyat_aktif) and pd.notna(med_fiyat_pasif) and med_fiyat_pasif > 0:
    fiyat_fark = (med_fiyat_aktif - med_fiyat_pasif) / med_fiyat_pasif * 100
    if fiyat_fark > 5:
        fiyat_trend = f"Aktif ilanların fiyatı, kapanan ilanların {abs(fiyat_fark):.0f}% üzerinde — fiyatlar yükseliyor."
    elif fiyat_fark < -5:
        fiyat_trend = f"Aktif ilanların fiyatı, kapanan ilanların {abs(fiyat_fark):.0f}% altında — fiyatlar geriliyor."
    else:
        fiyat_trend = f"Aktif ve kapanan ilan fiyatları birbirine yakın (%{abs(fiyat_fark):.0f} fark) — fiyatlar stabil."

sure_yorum = ""
if pd.notna(med_sure_a) and pd.notna(med_sure_p):
    if med_sure_a > med_sure_p * 1.3:
        sure_yorum = f"Aktif ilanlar ortalama {med_sure_a:.0f} gündür bekliyor, kapananlar {med_sure_p:.0f} günde kapanmıştı — piyasa yavaşlıyor."
    elif med_sure_a < med_sure_p * 0.7:
        sure_yorum = f"Aktif ilanlar ortalama {med_sure_a:.0f} gündür bekliyor, kapananlar {med_sure_p:.0f} günde kapanmıştı — piyasa hızlanıyor."
    else:
        sure_yorum = f"İlanlar ortalama {med_sure_p:.0f} günde kapanıyor, aktif ilanlar {med_sure_a:.0f} gündür bekliyor."

st.info(f"""**📊 Pazar Değerlendirmesi**

**{pazar_hiz}** — {pazar_tip}.

{hacim_notu}

{fiyat_trend} {sure_yorum}

Son {meta['donem']} içinde **{kapanan} ilan** kapandı \
(dönem başı {donem_bas_stok_n} stok üzerinden %{absorp:.0f} absorpsiyon). Şu an piyasada **{len(df_a)} aktif ilan** bekliyor.""" if donem_bas_stok_n > 0 else "")

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# BÖLÜM 2: AYLIK PAZAR AKIŞI
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("### 2️⃣ Aylık Pazar Akışı")
st.info("📌 **Metodoloji:** Dönem başı stok, 2024 başından itibaren çekilen verilerle hesaplanmaktadır. Her ayın dönem başı stoku = o aydan önce giren ve henüz kapanmamış tüm ilanlar. Analiz penceresi dışında kapanmış ilanlar dahil edilememektedir — gerçek stok bu değerin hafif üzerinde olabilir.")

grafik_listesi = {}
aylik = pd.DataFrame()  # Guard — tarih verisi yoksa boş kalır

df_tum_giris = pd.concat([df_a, df_p], ignore_index=True) if not df_a.empty and not df_p.empty else (df_a if not df_a.empty else df_p)

if "__ay" in df_tum_giris.columns:
    aylik_giren = df_tum_giris.groupby("__ay").size().reset_index(name="Giren")
    aylik_giren["Ay"] = aylik_giren["__ay"].astype(str)

    if not df_p.empty and "__kalk_ay" in df_p.columns:
        aylik_cikan = df_p.dropna(subset=["__kalk_ay"]).groupby("__kalk_ay").agg(
            Cikan=("__birim","count"),
            Med_Sure=("__sure","median"),
            Med_m2_Fiyat=("__birim","median"),
            Med_Fiyat=("__fiyat","median"),
        ).reset_index()
        aylik_cikan["Ay"] = aylik_cikan["__kalk_ay"].astype(str)
    else:
        aylik_cikan = pd.DataFrame(columns=["Ay","Cikan","Med_Sure","Med_m2_Fiyat","Med_Fiyat"])

    aylik = aylik_giren.merge(
        aylik_cikan[["Ay","Cikan","Med_Sure","Med_m2_Fiyat","Med_Fiyat"]] if not aylik_cikan.empty else pd.DataFrame(columns=["Ay","Cikan","Med_Sure","Med_m2_Fiyat","Med_Fiyat"]),
        on="Ay", how="outer"
    ).fillna(0).sort_values("Ay").reset_index(drop=True)

    aylik["Giren"] = aylik["Giren"].astype(int)
    aylik["Cikan"] = aylik["Cikan"].astype(int)
    aylik["Net"]   = aylik["Giren"] - aylik["Cikan"]

    stok_satirlar = []
    for _, r in aylik.iterrows():
        try:
            ay_ts = pd.Period(r["Ay"], freq="M").to_timestamp()
        except Exception:
            stok_satirlar.append((0, 0, None)); continue
        d_bas = donem_basi_stok(df_a, df_p, ay_ts)
        d_son = d_bas + int(r["Giren"]) - int(r["Cikan"])
        abs_pct = (int(r["Cikan"]) / d_bas * 100) if d_bas > 0 else None
        stok_satirlar.append((d_bas, d_son, abs_pct))

    aylik["Donem_Bas_Stok"]   = [s[0] for s in stok_satirlar]
    aylik["Donem_Son_Stok"]   = [s[1] for s in stok_satirlar]
    aylik["Aylik_Absorp_Pct"] = [s[2] for s in stok_satirlar]

    gun_sec = meta.get("gun", 365)
    sinir_ay = (datetime.now() - timedelta(days=gun_sec)).strftime("%Y-%m")
    aylik_goster = aylik[aylik["Ay"] >= sinir_ay].copy()

    son3 = aylik_goster.tail(3)
    sure_trend = son3["Med_Sure"].tolist()
    if len(sure_trend) >= 2 and sure_trend[-1] > sure_trend[0] * 1.3:
        st.warning(f"⚠️ **TREND ALARMI:** Medyan ilan süresi son {len(sure_trend)} ayda "
                   f"{sure_trend[0]:.0f} günden {sure_trend[-1]:.0f} güne yükseldi — piyasa yavaşlıyor.")

    gc1, gc2 = st.columns(2)
    with gc1:
        fig1 = go.Figure()
        fig1.add_bar(x=aylik_goster["Ay"], y=aylik_goster["Giren"], name="Yeni Giren", marker_color="#22c55e")
        fig1.add_bar(x=aylik_goster["Ay"], y=aylik_goster["Cikan"], name="Kapanan", marker_color="#ef4444")
        fig1.add_scatter(x=aylik_goster["Ay"], y=aylik_goster["Donem_Son_Stok"],
            name="Stok", mode="lines+markers",
            line=dict(color="#1e2d3d", width=1.5, dash="dot"),
            yaxis="y2", marker=dict(size=4))
        fig1.update_layout(
            title="Aylık Giriş / Kapanış + Stok",
            barmode="group", plot_bgcolor="white", paper_bgcolor="white",
            legend=dict(orientation="h", y=1.12), height=300,
            margin=dict(l=0,r=0,t=45,b=0), font=dict(size=10),
            yaxis=dict(title="Adet"),
            yaxis2=dict(title="Stok", overlaying="y", side="right", showgrid=False))
        st.plotly_chart(fig1, use_container_width=True)
        grafik_listesi["aylik_giris_cikis"] = fig1

    with gc2:
        _fiyat_ay = aylik_goster[aylik_goster["Med_m2_Fiyat"] > 0]
        if not _fiyat_ay.empty:
            fig2 = go.Figure()
            fig2.add_scatter(x=_fiyat_ay["Ay"], y=_fiyat_ay["Med_m2_Fiyat"],
                mode="lines+markers", name="Kapanış m² Fiyatı",
                line=dict(color="#1e2d3d", width=2), marker=dict(size=6))
            fig2.update_layout(
                title="Aylık Kapanış m² Fiyatı Trendi",
                plot_bgcolor="white", paper_bgcolor="white", height=300,
                margin=dict(l=0,r=0,t=45,b=0), font=dict(size=10))
            st.plotly_chart(fig2, use_container_width=True)
            grafik_listesi["aylik_fiyat_trend"] = fig2

    st.markdown("<br>", unsafe_allow_html=True)
    for _, r in aylik_goster.iterrows():
        giren = int(r.get("Giren",0)); cikan = int(r.get("Cikan",0)); net = int(r.get("Net",0))
        sure  = r.get("Med_Sure",0);  fiyat = r.get("Med_m2_Fiyat",0)
        d_bas = int(r.get("Donem_Bas_Stok",0)); d_son = int(r.get("Donem_Son_Stok",0))
        abs_p = r.get("Aylik_Absorp_Pct"); ay_str = str(r["Ay"])
        abs_badge = (f'<span style="background:{"#16a34a" if abs_p and abs_p>20 else "#d97706" if abs_p and abs_p>10 else "#dc2626"};'
                     f'color:white;padding:1px 6px;border-radius:4px;font-size:9px;font-weight:700;">%{abs_p:.0f} absorpsiyon</span>'
                     if abs_p is not None else "")
        net_yon   = (f"▲ +{net} stok artıyor" if net > 5 else f"▼ {net} stok azalıyor" if net < -5 else f"≈ {net:+d} denge")
        net_renk  = "#dc2626" if net > 5 else "#16a34a" if net < -5 else "#64748b"
        sure_yon  = ("⚡ çok hızlı" if sure > 0 and sure < 30 else "✅ normal" if sure > 0 and sure < 60 else "⏳ uzun" if sure > 0 else "")
        st.markdown(f"""
<div style="border:1px solid #e2e8f0;border-radius:8px;padding:10px 16px;margin-bottom:5px;background:white;">
  <div style="display:flex;gap:16px;align-items:center;flex-wrap:wrap;font-size:12px;">
    <div style="min-width:70px;font-weight:700;color:#1e2d3d;font-size:13px;">{ay_str}</div>
    <span>🟢 <b>{giren}</b> giren</span><span>🔴 <b>{cikan}</b> kapandı</span>
    <span style="color:{net_renk};font-weight:600;">{net_yon}</span>
    <span style="color:#475569;">📦 Stok: <b>{d_bas}</b>→<b>{d_son}</b></span>
    {abs_badge}
    {"<span>⏱ <b>" + str(int(sure)) + " gün</b> " + sure_yon + "</span>" if sure > 0 else ""}
    {"<span>💰 <b>" + _fmt(fiyat) + "/m²</b></span>" if fiyat > 0 else ""}
  </div>
</div>""", unsafe_allow_html=True)

    if len(aylik_goster) >= 2:
        son3 = aylik_goster.tail(3)
        t_giren = int(son3["Giren"].sum()); t_cikan = int(son3["Cikan"].sum()); t_net = t_giren - t_cikan
        ort_sure = son3[son3["Med_Sure"] > 0]["Med_Sure"].mean()
        avg_abs  = son3["Aylik_Absorp_Pct"].dropna().mean()
        stok_yorum  = (f"Stok **{t_net} ilan artıyor** — arz fazlası oluşuyor" if t_net > 0 else f"Stok **{abs(t_net)} ilan azalıyor** — talep güçlü")
        sure_yorum2 = (f"İlanlar ortalama **{ort_sure:.0f} günde** kapanıyor — " + ("çok hızlı piyasa." if ort_sure < 30 else "normal satış süresi." if ort_sure < 60 else "uzun bekleme, pazarlık payı var.") if pd.notna(ort_sure) and ort_sure > 0 else "")
        abs_yorum   = f"Aylık ortalama absorpsiyon **%{avg_abs:.1f}**." if pd.notna(avg_abs) else ""
        st.info(f"**Son 3 Ay:** {t_giren} ilan girdi, {t_cikan} kapandı. {stok_yorum}. {sure_yorum2} {abs_yorum}")
else:
    st.caption("Tarih verisi mevcut değil.")

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# BÖLÜM 3: FİYAT SEGMENTASYONU (DOM)
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("### 3️⃣ Fiyat Segmentasyonu — İlan Süresi (Aktif İlanlar)")

dom_grp    = pd.DataFrame()
_aktif_seg = pd.DataFrame()

if not df_a.empty and df_a["__sure"].notna().any():
    dom_bins   = [0, 30, 60, 90, 180, float("inf")]
    dom_labels = ["0-30 gün 🟢","31-60 gün 🟡","61-90 gün 🟠","91-180 gün 🔴","180+ gün ⚫"]
    df_a["__dom"] = pd.cut(df_a["__sure"], bins=dom_bins, labels=dom_labels, right=True)
    dom_grp = df_a.groupby("__dom", observed=True).agg(
        Adet=("__birim","count"), Med_m2=("__m2","median"),
        Med_Fiyat=("__fiyat","median"), Ort_Fiyat=("__fiyat","mean"),
        Med_m2_Fiyat=("__birim","median"), Ort_m2_Fiyat=("__birim","mean"),
    ).reset_index().rename(columns={"__dom":"İlan Süresi"})
    dom_grp["Med_Ref_100m2"] = dom_grp["Med_m2_Fiyat"] * 100
    dom_grp = dom_grp[dom_grp["Adet"] > 0]
    for c in ["Med_m2","Med_Fiyat","Ort_Fiyat","Med_m2_Fiyat","Ort_m2_Fiyat","Med_Ref_100m2"]:
        dom_grp[c] = dom_grp[c].round(0)

    dc1, dc2 = st.columns([1.5, 1])
    with dc1:
        st.dataframe(dom_grp, use_container_width=True, hide_index=True,
            column_config={
                "İlan Süresi": st.column_config.TextColumn("İlan Süresi"),
                "Adet": st.column_config.NumberColumn("Adet", format="%d"),
                "Med_m2": st.column_config.NumberColumn("Med. M²", format="%.0f m²"),
                "Med_Fiyat": st.column_config.NumberColumn("Med. Fiyat", format="%.0f ₺"),
                "Ort_Fiyat": st.column_config.NumberColumn("Ort. Fiyat", format="%.0f ₺"),
                "Med_m2_Fiyat": st.column_config.NumberColumn("Med. m² Fiyatı ★", format="%.0f ₺"),
                "Ort_m2_Fiyat": st.column_config.NumberColumn("Ort. m² Fiyatı", format="%.0f ₺"),
                "Med_Ref_100m2": st.column_config.NumberColumn("Med. 100m² Ref.", format="%.0f ₺"),
            })
    with dc2:
        fig3 = px.bar(dom_grp, x="Med_m2_Fiyat", y="İlan Süresi", orientation="h",
            text="Adet", color="Adet", color_continuous_scale=["#22c55e","#f59e0b","#ef4444"])
        fig3.update_layout(plot_bgcolor="white", paper_bgcolor="white",
            height=280, margin=dict(l=0,r=10,t=5,b=0), font=dict(size=10), coloraxis_showscale=False)
        fig3.update_traces(textposition="outside", textfont_size=10)
        st.plotly_chart(fig3, use_container_width=True)
        grafik_listesi["dom_segment"] = fig3

    _aktif_seg = dom_grp[dom_grp["İlan Süresi"] == "31-60 gün 🟡"]
    if not _aktif_seg.empty:
        _med = _aktif_seg.iloc[0]["Med_Ref_100m2"]
        _ort = _aktif_seg.iloc[0]["Ort_m2_Fiyat"] * 100
        st.info(f"**Aktif Satış Penceresi (31-60 gün):** {_aktif_seg.iloc[0]['Adet']} ilan — Medyan 100m² fiyatı **{_fmt(_med)}** / Ortalama **{_fmt(_ort)}**")
else:
    st.caption("İlan süresi verisi mevcut değil.")

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# BÖLÜM 4: ÜRÜN ANALİZİ
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("### 4️⃣ Ürün Analizi")

oda_grp = pd.DataFrame()
ua1, ua2 = st.columns(2)

with ua1:
    st.markdown("**Oda Tipine Göre (Aktif)**")
    if not df_a.empty and "Oda sayısı" in df_a.columns:
        oda_grp = df_a.groupby("Oda sayısı").agg(
            Adet=("__birim","count"), Med_m2=("__m2","median"),
            Med_Fiyat=("__fiyat","median"), Med_m2_Fiyat=("__birim","median"),
            Med_Sure=("__sure","median"),
        ).reset_index().sort_values("Adet", ascending=False).head(10)
        for c in ["Med_m2","Med_Fiyat","Med_m2_Fiyat","Med_Sure"]:
            oda_grp[c] = oda_grp[c].round(0)
        st.dataframe(oda_grp, use_container_width=True, hide_index=True,
            column_config={
                "Adet": st.column_config.NumberColumn("Adet", format="%d"),
                "Med_m2": st.column_config.NumberColumn("Med. M²", format="%.0f"),
                "Med_Fiyat": st.column_config.NumberColumn("Med. Fiyat", format="%.0f ₺"),
                "Med_m2_Fiyat": st.column_config.NumberColumn("Med. m² Fiyatı", format="%.0f ₺"),
                "Med_Sure": st.column_config.NumberColumn("Med. Süre", format="%.0f gün"),
            })

with ua2:
    st.markdown("**Yeni / Eski Bina Analizi (Aktif + Pasif)**")
    df_tum = pd.concat([df_a, df_p], ignore_index=True) if not df_a.empty and not df_p.empty else (df_a if not df_a.empty else df_p)
    if "Bina Yaşı" in df_tum.columns:
        # ── PATCH 1: Bina yaşı mapping — Revy gerçek formatına göre ─────────
        # Yeni binalar tekil sayı (0,1,2,3,4,5), diğerleri aralık (6-10, 11-15...)
        YAS_MAP = {
            "Yeni (0-5 yıl)":   ["0","1","2","3","4","5"],
            "Genç (6-15 yıl)":  ["6-10","11-15"],
            "Orta (16-25 yıl)": ["16-20","21-25"],
            "Eski (26+ yıl)":   ["26-30","30 üstü"],
        }
        def yas_grup(v):
            s = str(v).strip()
            for g, vals in YAS_MAP.items():
                if s in vals: return g
            return None  # sayısal fallback yok — map tüm Revy değerlerini kapsıyor

        df_tum["__yas"] = df_tum["Bina Yaşı"].apply(yas_grup)
        yas_grp = df_tum[df_tum["__yas"].notna()].groupby("__yas").agg(
            Adet=("__birim","count"),
            Med_m2_Fiyat=("__birim","median"),
            Ort_m2_Fiyat=("__birim","mean"),
        ).reset_index().rename(columns={"__yas":"Bina Yaşı Grubu"})
        yas_grp["Med_Ref"] = yas_grp["Med_m2_Fiyat"] * 100
        yas_sirasi = list(YAS_MAP.keys())
        yas_grp["_s"] = yas_grp["Bina Yaşı Grubu"].map({v:i for i,v in enumerate(yas_sirasi)})
        yas_grp = yas_grp.sort_values("_s").drop(columns=["_s"])
        for c in ["Med_m2_Fiyat","Ort_m2_Fiyat","Med_Ref"]:
            yas_grp[c] = yas_grp[c].round(0)

        st.dataframe(yas_grp, use_container_width=True, hide_index=True,
            column_config={
                "Adet": st.column_config.NumberColumn("Adet", format="%d"),
                "Med_m2_Fiyat": st.column_config.NumberColumn("Med. m² Fiyatı ★", format="%.0f ₺"),
                "Ort_m2_Fiyat": st.column_config.NumberColumn("Ort. m² Fiyatı", format="%.0f ₺"),
                "Med_Ref": st.column_config.NumberColumn("Med. 100m² Ref.", format="%.0f ₺"),
            })

        _y = yas_grp[yas_grp["Bina Yaşı Grubu"]=="Yeni (0-5 yıl)"]
        _e = yas_grp[yas_grp["Bina Yaşı Grubu"]=="Eski (26+ yıl)"]
        if not _y.empty and not _e.empty:
            _prem = (_y.iloc[0]["Med_m2_Fiyat"] - _e.iloc[0]["Med_m2_Fiyat"]) / _e.iloc[0]["Med_m2_Fiyat"] * 100
            st.info(f"**Yeni/Eski Bina Farkı:** Yeni binalar eski binalara göre m² bazında **%{abs(_prem):.1f} {'daha pahalı' if _prem > 0 else 'daha ucuz'}**")

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# BÖLÜM 4.5: MAHALLE KIRILIMI
# ─────────────────────────────────────────────────────────────────────────────
if not df_a.empty and "Mahalle" in df_a.columns and df_a["Mahalle"].nunique() > 1:
    st.markdown("### 🗺 Mahalle Kırılımı (Aktif İlanlar)")
    mah_grp = df_a.groupby("Mahalle").agg(
        Adet=("__birim","count"), Med_m2_Fiyat=("__birim","median"),
        Med_Fiyat=("__fiyat","median"), Med_Sure=("__sure","median"),
    ).reset_index().sort_values("Adet", ascending=False).head(12)
    mah_grp["Med_Ref_100m2"] = mah_grp["Med_m2_Fiyat"] * 100

    if not df_p.empty and "Mahalle" in df_p.columns:
        mah_p = df_p.groupby("Mahalle").agg(Kapanan=("__birim","count")).reset_index()
        mah_grp = mah_grp.merge(mah_p, on="Mahalle", how="left").fillna(0)
        mah_grp["Kapanan"] = mah_grp["Kapanan"].astype(int)
        # Not: closure rate (kapanan/toplam) — gerçek absorpsiyon değil, referans
        mah_grp["Kapanan_Oran"] = (mah_grp["Kapanan"] / (mah_grp["Adet"] + mah_grp["Kapanan"]) * 100).round(1)

    for c in ["Med_m2_Fiyat","Med_Fiyat","Med_Sure","Med_Ref_100m2"]:
        mah_grp[c] = mah_grp[c].round(0)

    col_cfg_mah = {
        "Adet": st.column_config.NumberColumn("Aktif", format="%d"),
        "Med_m2_Fiyat": st.column_config.NumberColumn("Med. m² Fiyatı ★", format="%.0f ₺"),
        "Med_Fiyat": st.column_config.NumberColumn("Med. Fiyat", format="%.0f ₺"),
        "Med_Sure": st.column_config.NumberColumn("Bekleme Süresi", format="%.0f gün"),
        "Med_Ref_100m2": st.column_config.NumberColumn("100m² Ref.", format="%.0f ₺"),
    }
    if "Kapanan" in mah_grp.columns:
        col_cfg_mah["Kapanan"] = st.column_config.NumberColumn("Kapanan", format="%d")
        col_cfg_mah["Kapanan_Oran"] = st.column_config.NumberColumn("Kapanış Oranı%*", format="%.1f%%")

    st.dataframe(mah_grp, use_container_width=True, hide_index=True, column_config=col_cfg_mah)
    st.caption("* Kapanış Oranı = kapanan / (aktif + kapanan) — mahalle bazlı dönem başı stok olmadığından gerçek absorpsiyon değildir.")

    if len(mah_grp) >= 2:
        en_hizli  = mah_grp.loc[mah_grp["Med_Sure"].idxmin()]
        en_pahali = mah_grp.loc[mah_grp["Med_m2_Fiyat"].idxmax()]
        st.caption(f"⚡ En hızlı: **{en_hizli['Mahalle']}** ({en_hizli['Med_Sure']:.0f} gün) — 💰 En pahalı: **{en_pahali['Mahalle']}** ({_fmt(en_pahali['Med_m2_Fiyat'])}/m²)")

    st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# BÖLÜM 5: PAZAR OYUNCULARI
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("### 5️⃣ Pazar Oyuncuları")

ofis_grp = pd.DataFrame()  # scope güvencesi

MARKA_RENK = {
    "startkey":"#1e2d3d","remax":"#E85D75","turpa":"#F59E0B",
    "turyap":"#8B5CF6","coldwell":"#06B6D4","kw":"#10B981",
    "alesta":"#F97316","viya":"#EC4899","orsa":"#84CC16",
    "century21":"#6366F1","bagimsiz":"#94A3B8","mulk_sahibi":"#CBD5E1",
}
def marka_etiket(m):
    _MAP = {"bagimsiz":"Yerel Ofis","mulk_sahibi":"Mülk Sahibi"}
    return _MAP.get(str(m).lower(), str(m).upper()) if m else "-"

po1, po2 = st.columns([1.5, 1])
with po1:
    st.markdown("**Ofis Bazlı Pazar Payı (Aktif, İlk 15)**")
    if not df_a.empty and "__ofis" in df_a.columns and "MARKA" in df_a.columns:
        df_a["__ofis_g"] = df_a["__ofis"].fillna("Mülk Sahibi")
        ofis_grp = df_a.groupby(["__ofis_g","MARKA"]).agg(
            Adet=("__birim","count"), Med_m2_Fiyat=("__birim","median"), Med_Sure=("__sure","median"),
        ).reset_index().sort_values("Adet", ascending=False).head(15)
        ofis_grp["Pay"] = (ofis_grp["Adet"] / len(df_a) * 100).round(1)
        ofis_grp["MARKA_ETİ"] = ofis_grp["MARKA"].apply(marka_etiket)
        for c in ["Med_m2_Fiyat","Med_Sure"]: ofis_grp[c] = ofis_grp[c].round(0)
        fig4 = px.bar(ofis_grp.sort_values("Adet"), x="Adet", y="__ofis_g",
            orientation="h", color="MARKA",
            color_discrete_map={r["MARKA"]: MARKA_RENK.get(r["MARKA"],"#94A3B8") for _,r in ofis_grp.iterrows()},
            text="Adet")
        fig4.update_layout(showlegend=False, plot_bgcolor="white", paper_bgcolor="white",
            height=420, margin=dict(l=0,r=10,t=5,b=0), font=dict(size=9), yaxis=dict(title=""))
        fig4.update_traces(textposition="outside", textfont_size=9)
        st.plotly_chart(fig4, use_container_width=True)
        grafik_listesi["ofis_pazar_payi"] = fig4

with po2:
    st.markdown("**Marka Dağılımı (Aktif)**")
    if not df_a.empty and "MARKA" in df_a.columns:
        marka_grp = df_a.groupby("MARKA").agg(
            Adet=("__birim","count"), Med_m2_Fiyat=("__birim","median"),
        ).reset_index().sort_values("Adet", ascending=False)
        marka_grp["Pay"] = (marka_grp["Adet"] / len(df_a) * 100).round(1)
        marka_grp["Etiket"] = marka_grp["MARKA"].apply(marka_etiket)
        marka_grp["Med_m2_Fiyat"] = marka_grp["Med_m2_Fiyat"].round(0)
        fig5 = px.pie(marka_grp, values="Adet", names="Etiket", color="MARKA",
            color_discrete_map={r["MARKA"]: MARKA_RENK.get(r["MARKA"],"#94A3B8") for _,r in marka_grp.iterrows()},
            hole=0.4)
        fig5.update_layout(height=300, margin=dict(l=0,r=0,t=20,b=0), font=dict(size=10),
            legend=dict(orientation="v", x=1, y=0.5))
        fig5.update_traces(textposition="inside", textinfo="percent+label", textfont_size=9)
        st.plotly_chart(fig5, use_container_width=True)
        grafik_listesi["marka_dagilimi"] = fig5

    st.markdown("**GD Bazlı (Aktif, İlk 10)**")
    if not df_a.empty and "İlan sahibi" in df_a.columns:
        gd_cols = ["İlan sahibi","__birim","__sure"]
        if "__ofis" in df_a.columns: gd_cols.append("__ofis")
        gd_grp = df_a[gd_cols].copy()
        if "__ofis" in gd_grp.columns:
            gd_grp = gd_grp.groupby(["İlan sahibi","__ofis"]).agg(
                Adet=("__birim","count"), Med_m2_Fiyat=("__birim","median"), Med_Sure=("__sure","median"),
            ).reset_index().sort_values("Adet", ascending=False).head(10)
            gd_grp.rename(columns={"__ofis":"Ofis"}, inplace=True)
        else:
            gd_grp = gd_grp.groupby("İlan sahibi").agg(
                Adet=("__birim","count"), Med_m2_Fiyat=("__birim","median"), Med_Sure=("__sure","median"),
            ).reset_index().sort_values("Adet", ascending=False).head(10)
        gd_grp["Med_m2_Fiyat"] = gd_grp["Med_m2_Fiyat"].round(0)
        gd_grp["Med_Sure"] = gd_grp["Med_Sure"].round(0)
        col_cfg = {
            "İlan sahibi": st.column_config.TextColumn("Danışman"),
            "Adet": st.column_config.NumberColumn("İlan", format="%d"),
            "Med_m2_Fiyat": st.column_config.NumberColumn("m² Fiyatı", format="%.0f ₺"),
            "Med_Sure": st.column_config.NumberColumn("Bekleme Süresi", format="%.0f gün"),
        }
        if "Ofis" in gd_grp.columns: col_cfg["Ofis"] = st.column_config.TextColumn("Ofis")
        st.dataframe(gd_grp, use_container_width=True, hide_index=True, column_config=col_cfg)

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# BÖLÜM 6: AYKIRI DEĞER & MÜKERRER ANALİZİ
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("### 6️⃣ Aykırı Değer & Mükerrer Analizi")

av1, av2 = st.columns(2)
with av1:
    if not df_a.empty:
        if all(c in df_a.columns for c in ["Fiyat","M2","Mahalle"]):
            df_a["_dk"] = df_a["Fiyat"].astype(str)+"_"+df_a["M2"].astype(str)+"_"+df_a["Mahalle"].astype(str)
            dupe_n = df_a["_dk"].duplicated(keep=False).sum()
            st.metric("Mükerrer İlan (Aktif)", f"{dupe_n}", f"%{dupe_n/len(df_a)*100:.1f}")
        med_b = df_a["__birim"].median(); ort_b = df_a["__birim"].mean()
        if pd.notna(med_b) and pd.notna(ort_b) and ort_b > 0:
            fark = abs(ort_b - med_b) / ort_b * 100
            st.metric("Med./Ort. m² Fiyatı Farkı", f"{_fmt(med_b)} / {_fmt(ort_b)}",
                f"%{fark:.1f} {'⚠️ aykırı var' if fark > 10 else '✅ dengeli'}")

with av2:
    if not df_a.empty and df_a["__sure"].notna().any():
        st.markdown("**En uzun süre bekleyen 5 ilan**")
        _uzun_cols = [c for c in ["İlan Başlığı","__sure","__birim","Ofis"] if c in df_a.columns]
        uzun = df_a.nlargest(5, "__sure")[_uzun_cols].copy()
        uzun.columns = ["Başlık","Gün","m² Fiyatı","Ofis"][:len(_uzun_cols)]
        uzun["Gün"] = uzun["Gün"].round(0).astype(int)
        if "m² Fiyatı" in uzun.columns: uzun["m² Fiyatı"] = uzun["m² Fiyatı"].round(0)
        st.dataframe(uzun, use_container_width=True, hide_index=True,
            column_config={
                "Gün": st.column_config.NumberColumn("Gün", format="%d"),
                "m² Fiyatı": st.column_config.NumberColumn("m² Fiyatı", format="%.0f ₺"),
            })

st.divider()

# ─────────────────────────────────────────────────────────────────────────────
# RAPOR ÇIKTISI — PDF + WORD
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("### 📥 Rapor Çıktısı")

dl1, dl2, dl3 = st.columns(3)

with dl1:
    if st.button("📄 PDF Raporu Oluştur", use_container_width=True, key="rpr_pdf"):
        with st.spinner("PDF oluşturuluyor..."):
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib import colors
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import cm
                from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                    Table, TableStyle, Image as RLImage, PageBreak, HRFlowable)
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                import io as _io

                try:
                    pdfmetrics.registerFont(TTFont("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
                    pdfmetrics.registerFont(TTFont("DejaVu-Bold", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"))
                    FONT = "DejaVu"; FONT_B = "DejaVu-Bold"
                except Exception:
                    FONT = "Helvetica"; FONT_B = "Helvetica-Bold"

                buf = _io.BytesIO()
                doc = SimpleDocTemplate(buf, pagesize=A4,
                    leftMargin=1.5*cm, rightMargin=1.5*cm, topMargin=2*cm, bottomMargin=2*cm)

                s_title = ParagraphStyle("title", fontName=FONT_B, fontSize=18, textColor=colors.HexColor("#1e2d3d"), spaceAfter=6)
                s_sub   = ParagraphStyle("sub",   fontName=FONT,   fontSize=10, textColor=colors.HexColor("#64748b"), spaceAfter=12)
                s_h1    = ParagraphStyle("h1",    fontName=FONT_B, fontSize=13, textColor=colors.HexColor("#1e2d3d"), spaceBefore=16, spaceAfter=8, borderPad=4, backColor=colors.HexColor("#f1f5f9"), leftIndent=6)
                s_info  = ParagraphStyle("info",  fontName=FONT,   fontSize=9,  textColor=colors.HexColor("#1e40af"), spaceAfter=8, leading=13, backColor=colors.HexColor("#eff6ff"), borderPad=6)

                def tablo(data, col_widths=None, header_bg="#1e2d3d"):
                    style = TableStyle([
                        ("BACKGROUND", (0,0), (-1,0), colors.HexColor(header_bg)),
                        ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
                        ("FONTNAME",   (0,0), (-1,0), FONT_B),
                        ("FONTSIZE",   (0,0), (-1,-1), 8),
                        ("FONTNAME",   (0,1), (-1,-1), FONT),
                        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f8fafc")]),
                        ("GRID",       (0,0), (-1,-1), 0.3, colors.HexColor("#dce4ee")),
                        ("ALIGN",      (1,0), (-1,-1), "RIGHT"),
                        ("ALIGN",      (0,0), (0,-1),  "LEFT"),
                        ("TOPPADDING",    (0,0), (-1,-1), 4),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                        ("LEFTPADDING",   (0,0), (-1,-1), 5),
                        ("RIGHTPADDING",  (0,0), (-1,-1), 5),
                    ])
                    t = Table(data, colWidths=col_widths, repeatRows=1)
                    t.setStyle(style)
                    return t

                story = []
                story.append(Spacer(1, 1*cm))
                story.append(Paragraph("PAZAR ANALİZ RAPORU", s_title))
                story.append(Paragraph(baslik, ParagraphStyle("bl", fontName=FONT_B, fontSize=12, textColor=colors.HexColor("#374151"), spaceAfter=4)))
                story.append(Paragraph(f"Rapor Tarihi: {meta['tarih']} &nbsp;|&nbsp; Aktif: {len(df_a)} ilan &nbsp;|&nbsp; Pasif ({meta['donem']}): {len(df_p)} ilan", s_sub))
                story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1e2d3d")))
                story.append(Spacer(1, 0.5*cm))

                # 1. Pazar Özeti
                story.append(Paragraph("1. Pazar Ozeti", s_h1))
                ozet_data = [
                    ["Metrik","Aktif","Pasif"],
                    ["Ilan Sayisi", str(len(df_a)), str(len(df_p))],
                    ["Absorpsiyon Orani", f"%{absorp:.0f} ({pazar_hiz})", ""],
                    ["Medyan m2 Fiyati", _fmt(df_a["__birim"].median()), _fmt(df_p["__birim"].median())],
                    ["Medyan Fiyat",     _fmt(df_a["__fiyat"].median()), _fmt(df_p["__fiyat"].median())],
                    ["Medyan Ilan Suresi", _fmt(df_a["__sure"].median(),"gun"), _fmt(df_p["__sure"].median(),"gun")],
                    ["FSBO Ilan (Mulk Sahibi)", str(fsbo_a), str(fsbo_p)],
                ]
                story.append(tablo(ozet_data, col_widths=[8*cm,5*cm,5*cm]))
                story.append(Spacer(1, 0.4*cm))

                # 2. Aylık Akış
                if not aylik.empty:
                    story.append(Paragraph("2. Aylik Pazar Akisi", s_h1))
                    _sinir_ay  = (datetime.now() - timedelta(days=meta.get("gun",365))).strftime("%Y-%m")
                    _aylik_pdf = aylik[aylik["Ay"] >= _sinir_ay]
                    ay_data = [["Ay","Giren","Kapanan","Net","D.Bas","D.Son","Abs%","Sure","Kap.m2"]]
                    for _, r in _aylik_pdf.iterrows():
                        abs_p = r.get("Aylik_Absorp_Pct"); sure = r.get("Med_Sure",0); fiyat = r.get("Med_m2_Fiyat",0)
                        ay_data.append([str(r["Ay"]), str(int(r.get("Giren",0))), str(int(r.get("Cikan",0))),
                            f"{int(r.get('Net',0)):+d}", str(int(r.get("Donem_Bas_Stok",0))), str(int(r.get("Donem_Son_Stok",0))),
                            f"%{abs_p:.0f}" if abs_p is not None else "-",
                            f"{sure:.0f}g" if sure > 0 else "-",
                            _fmt(fiyat) if fiyat > 0 else "-"])
                    story.append(tablo(ay_data, col_widths=[2.2*cm,1.2*cm,1.5*cm,1.2*cm,1.5*cm,1.5*cm,1.3*cm,1.3*cm,4.3*cm]))
                    story.append(Spacer(1, 0.3*cm))
                    for _g_key, _g_w, _g_h in [("aylik_giris_cikis",16*cm,6*cm),("aylik_fiyat_trend",16*cm,5.5*cm)]:
                        if _g_key in grafik_listesi:
                            try:
                                _img_b = pio.to_image(grafik_listesi[_g_key], format="png", width=700, height=260, scale=2)
                                story.append(RLImage(_io.BytesIO(_img_b), width=_g_w, height=_g_h))
                                story.append(Spacer(1, 0.2*cm))
                            except Exception: pass
                    story.append(Spacer(1, 0.3*cm))

                # 3. DOM
                if not dom_grp.empty:
                    story.append(Paragraph("3. Fiyat Segmentasyonu (Ilan Suresi)", s_h1))
                    dom_data = [["Ilan Suresi","Adet","Med. m2","Med. Fiyat","Med. m2 Fiyati","Med. 100m2 Ref."]]
                    for _, r in dom_grp.iterrows():
                        s_label = str(r["İlan Süresi"]).replace("🟢","").replace("🟡","").replace("🟠","").replace("🔴","").replace("⚫","").strip()
                        dom_data.append([s_label, str(int(r["Adet"])), f"{r['Med_m2']:.0f} m2",
                            _fmt(r["Med_Fiyat"]), _fmt(r["Med_m2_Fiyat"]), _fmt(r["Med_Ref_100m2"])])
                    story.append(tablo(dom_data, col_widths=[3*cm,1.5*cm,2.5*cm,4*cm,4*cm,3*cm]))
                    if not _aktif_seg.empty:
                        story.append(Spacer(1, 0.2*cm))
                        story.append(Paragraph(
                            f"Aktif Satis Penceresi (31-60 gun): {_aktif_seg.iloc[0]['Adet']} ilan — Medyan 100m2: {_fmt(_aktif_seg.iloc[0]['Med_Ref_100m2'])}",
                            s_info))
                    if "dom_segment" in grafik_listesi:
                        try:
                            _img_b = pio.to_image(grafik_listesi["dom_segment"], format="png", width=600, height=220, scale=2)
                            story.append(RLImage(_io.BytesIO(_img_b), width=14*cm, height=5*cm))
                        except Exception: pass

                # 4. Ürün
                if not oda_grp.empty:
                    story.append(Paragraph("4. Urun Analizi — Oda Tipi", s_h1))
                    oda_data = [["Oda","Adet","Med. m2","Med. Fiyat","Med. m2 Fiyati","Med. Sure"]]
                    for _, r in oda_grp.iterrows():
                        oda_data.append([str(r["Oda sayısı"]), str(int(r["Adet"])),
                            f"{r['Med_m2']:.0f}", _fmt(r["Med_Fiyat"]), _fmt(r["Med_m2_Fiyat"]), f"{r['Med_Sure']:.0f} gun"])
                    story.append(tablo(oda_data, col_widths=[2*cm,2*cm,2.5*cm,4*cm,4*cm,3.5*cm]))

                # 5. Pazar Oyuncuları
                story.append(PageBreak())
                story.append(Paragraph("5. Pazar Oyunculari", s_h1))
                if not ofis_grp.empty:
                    ofis_data = [["Ofis","Marka","Adet","Pay(%)","Med. m2 Fiyati","Med. Sure"]]
                    for _, r in ofis_grp.head(15).iterrows():
                        ofis_data.append([str(r["__ofis_g"])[:30], marka_etiket(r["MARKA"]),
                            str(int(r["Adet"])), f"%{r['Pay']:.1f}", _fmt(r["Med_m2_Fiyat"]), f"{r['Med_Sure']:.0f} gun"])
                    story.append(tablo(ofis_data, col_widths=[5*cm,2.5*cm,2*cm,2*cm,3.5*cm,3*cm]))
                    if "ofis_pazar_payi" in grafik_listesi:
                        try:
                            _img_b = pio.to_image(grafik_listesi["ofis_pazar_payi"], format="png", width=700, height=350, scale=2)
                            story.append(RLImage(_io.BytesIO(_img_b), width=16*cm, height=8*cm))
                        except Exception: pass

                story.append(Spacer(1, 1*cm))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dce4ee")))
                story.append(Paragraph(
                    f"Bu rapor Karma App / Pazar Radar tarafindan {meta['tarih']} tarihinde Revy.com.tr verileri kullanilarak olusturulmustur.",
                    ParagraphStyle("footer", fontName=FONT, fontSize=7, textColor=colors.HexColor("#94a3b8"), alignment=1, spaceBefore=6)))

                doc.build(story)
                buf.seek(0)
                dosya_adi = f"pazar_raporu_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
                st.download_button("⬇️ PDF İndir", data=buf.getvalue(),
                    file_name=dosya_adi, mime="application/pdf",
                    use_container_width=True, key="rpr_pdf_dl")
                st.success("✅ PDF hazır!")
            except Exception as e:
                st.error(f"PDF hatası: {e}")

with dl2:
    if st.button("📝 Word Raporu Oluştur", use_container_width=True, key="rpr_word"):
        with st.spinner("Word oluşturuluyor..."):
            try:
                from docx import Document as DocxDoc
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                import io as _io

                doc = DocxDoc()
                doc.styles["Normal"].font.name = "Arial"
                doc.styles["Normal"].font.size = Pt(10)

                def ekle_baslik(doc, metin, seviye=1):
                    p = doc.add_paragraph(); p.style = f"Heading {seviye}"
                    run = p.add_run(metin); run.font.name = "Arial"; return p

                def ekle_paragraf(doc, metin, bold=False, color=None, size=10):
                    p = doc.add_paragraph(); run = p.add_run(metin)
                    run.font.name = "Arial"; run.font.size = Pt(size); run.bold = bold
                    if color: run.font.color.rgb = RGBColor(*color)
                    return p

                def ekle_tablo(doc, data, header_color=(30,45,61)):
                    t = doc.add_table(rows=len(data), cols=len(data[0]))
                    t.style = "Table Grid"
                    for i, row_data in enumerate(data):
                        for j, cell_data in enumerate(row_data):
                            cell = t.rows[i].cells[j]
                            cell.text = str(cell_data)
                            run = cell.paragraphs[0].runs[0]
                            run.font.name = "Arial"; run.font.size = Pt(8)
                            if i == 0:
                                run.bold = True
                                run.font.color.rgb = RGBColor(255,255,255)
                                tcPr = cell._tc.get_or_add_tcPr()
                                shd = OxmlElement("w:shd")
                                shd.set(qn("w:fill"), "%02x%02x%02x" % header_color)
                                shd.set(qn("w:val"), "clear")
                                tcPr.append(shd)
                    return t

                p = doc.add_paragraph()
                run = p.add_run("PAZAR ANALİZ RAPORU")
                run.font.name = "Arial"; run.font.size = Pt(18); run.bold = True
                run.font.color.rgb = RGBColor(30,45,61)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                ekle_paragraf(doc, baslik, bold=True, size=12)
                ekle_paragraf(doc, f"Rapor Tarihi: {meta['tarih']}  |  Aktif: {len(df_a)} ilan  |  Pasif ({meta['donem']}): {len(df_p)} ilan", color=(100,116,139), size=9)
                doc.add_paragraph()

                # 1. Özet
                ekle_baslik(doc, "1. Pazar Özeti", 1)
                ekle_tablo(doc, [
                    ["Metrik","Aktif","Pasif"],
                    ["İlan Sayısı", str(len(df_a)), str(len(df_p))],
                    ["Absorpsiyon Oranı", f"%{absorp:.0f} ({pazar_hiz})", ""],
                    ["Medyan m² Fiyatı", _fmt(df_a["__birim"].median()), _fmt(df_p["__birim"].median())],
                    ["Medyan Fiyat",     _fmt(df_a["__fiyat"].median()), _fmt(df_p["__fiyat"].median())],
                    ["Medyan İlan Süresi", _fmt(df_a["__sure"].median(),"gün"), _fmt(df_p["__sure"].median(),"gün")],
                    ["FSBO İlan (Mülk Sahibi)", str(fsbo_a), str(fsbo_p)],
                ])
                doc.add_paragraph()

                # 2. Aylık
                if not aylik.empty:
                    ekle_baslik(doc, "2. Aylık Pazar Akışı", 1)
                    ay_data = [["Ay","Yeni Giren","Kapanan","Net","Ort. Süre","Kapanış m² Fiyatı"]]
                    for _, r in aylik.iterrows():
                        giren = int(r.get("Giren",0)); cikan = int(r.get("Cikan",0))
                        sure = r.get("Med_Sure",0); fiyat = r.get("Med_m2_Fiyat",0)
                        ay_data.append([str(r["Ay"]), str(giren), str(cikan),
                            f"{giren-cikan:+d}", f"{sure:.0f} gün" if sure > 0 else "-",
                            _fmt(fiyat) if fiyat > 0 else "-"])
                    ekle_tablo(doc, ay_data)
                    doc.add_paragraph()

                # 3. DOM
                if not dom_grp.empty:
                    ekle_baslik(doc, "3. Fiyat Segmentasyonu (İlan Süresi)", 1)
                    dom_data = [["İlan Süresi","Adet","Med. m² Fiyatı","Med. 100m² Ref."]]
                    for _, r in dom_grp.iterrows():
                        s_label = str(r["İlan Süresi"]).replace("🟢","").replace("🟡","").replace("🟠","").replace("🔴","").replace("⚫","").strip()
                        dom_data.append([s_label, str(int(r["Adet"])), _fmt(r["Med_m2_Fiyat"]), _fmt(r["Med_Ref_100m2"])])
                    ekle_tablo(doc, dom_data)
                    doc.add_paragraph()

                # 4. Ürün
                if not oda_grp.empty:
                    ekle_baslik(doc, "4. Ürün Analizi — Oda Tipi", 1)
                    oda_data = [["Oda","Adet","Med. m²","Med. Fiyat","Med. m² Fiyatı","Med. Süre"]]
                    for _, r in oda_grp.iterrows():
                        oda_data.append([str(r["Oda sayısı"]), str(int(r["Adet"])),
                            f"{r['Med_m2']:.0f}", _fmt(r["Med_Fiyat"]), _fmt(r["Med_m2_Fiyat"]), f"{r['Med_Sure']:.0f} gün"])
                    ekle_tablo(doc, oda_data)
                    doc.add_paragraph()

                # 5. Pazar Oyuncuları
                if not ofis_grp.empty:
                    ekle_baslik(doc, "5. Pazar Oyuncuları — Ofis Bazlı", 1)
                    of_data = [["Ofis","Marka","Adet","Pay (%)","Med. m² Fiyatı"]]
                    for _, r in ofis_grp.head(15).iterrows():
                        of_data.append([str(r["__ofis_g"])[:35], marka_etiket(r["MARKA"]),
                            str(int(r["Adet"])), f"%{r['Pay']:.1f}", _fmt(r["Med_m2_Fiyat"])])
                    ekle_tablo(doc, of_data)

                doc.add_paragraph()
                ekle_paragraf(doc, f"Bu rapor Karma App / Pazar Radar tarafından {meta['tarih']} tarihinde Revy.com.tr verileri kullanılarak oluşturulmuştur.", color=(148,163,184), size=8)

                buf = _io.BytesIO()
                doc.save(buf); buf.seek(0)
                dosya_adi = f"pazar_raporu_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                st.download_button("⬇️ Word İndir", data=buf.getvalue(),
                    file_name=dosya_adi,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key="rpr_word_dl")
                st.success("✅ Word hazır!")
            except Exception as e:
                st.error(f"Word hatası: {e}")

with dl3:
    if st.button("📊 Excel Raporu Oluştur", use_container_width=True, key="rpr_excel"):
        with st.spinner("Excel oluşturuluyor..."):
            try:
                import io as _io

                df_a_x = df_a.copy()
                df_a_x["Durum"] = "Aktif"
                df_p_x = df_p.copy()
                df_p_x["Durum"] = "Pasif"
                df_excel = pd.concat([df_a_x, df_p_x], ignore_index=True, sort=False)

                # İç hesaplama sütunlarını okunabilir isimlere çevir
                rename_map = {
                    "__birim": "m² Fiyatı",
                    "__fiyat": "Fiyat",
                    "__m2": "M²",
                    "__sure": "İlan Süresi (gün)",
                    "__ay": "Giriş Ayı",
                    "__kalk_ay": "Kalkış Ayı",
                    "__ofis": "Ofis",
                    "__ofis_g": "Ofis (Düzeltilmiş)",
                }
                df_excel = df_excel.rename(columns={k: v for k, v in rename_map.items() if k in df_excel.columns})

                # Kalan yardımcı/geçici sütunları (__dom, __yas, _dk vb.) çıkar
                _drop = [c for c in df_excel.columns if c.startswith("__") or c.startswith("_")]
                df_excel = df_excel.drop(columns=_drop, errors="ignore")

                # Durum en başa, ardından okunabilir sıralama
                _cols = ["Durum"] + [c for c in df_excel.columns if c != "Durum"]
                df_excel = df_excel[_cols]

                buf = _io.BytesIO()
                with pd.ExcelWriter(buf, engine="openpyxl") as writer:
                    df_excel.to_excel(writer, index=False, sheet_name="İlanlar")
                    ws = writer.sheets["İlanlar"]
                    ws.freeze_panes = "A2"
                    ws.auto_filter.ref = ws.dimensions
                    for i, col in enumerate(df_excel.columns, start=1):
                        try:
                            maxlen = max(df_excel[col].astype(str).map(len).max(), len(str(col))) + 2
                        except Exception:
                            maxlen = len(str(col)) + 2
                        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = min(maxlen, 42)
                buf.seek(0)
                dosya_adi = f"pazar_ilanlari_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
                st.download_button("⬇️ Excel İndir", data=buf.getvalue(),
                    file_name=dosya_adi,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True, key="rpr_excel_dl")
                st.success(f"✅ Excel hazır! ({len(df_excel)} ilan — {len(df_a)} aktif, {len(df_p)} pasif)")
            except Exception as e:
                st.error(f"Excel hatası: {e}")
